#!/usr/bin/env python3
"""
Convert a Chinese resume Markdown file to a styled Word (.docx) document.

Visual style mirrors markdown_resume_to_pdf.py:
gray bar section headings with left accent, name underline,
italic project intros, photo top-right, 微软雅黑 throughout.
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Reuse parsing helpers from the PDF exporter in the same directory.
_SCRIPT_DIR = Path(__file__).resolve().parent
if str(_SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPT_DIR))

from markdown_resume_to_pdf import (  # noqa: E402
    BULLET_TASK_RE,
    DATE_END_PATTERN,
    PRESETS,
    match_project_task,
    match_sub_item,
    parse_bold_project_header,
    parse_project_group,
    parse_row_heading,
)

FONT_NAME = "微软雅黑"
COLOR_BODY = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_CONTACT = RGBColor(0x55, 0x55, 0x55)
COLOR_PROJECT_INTRO = RGBColor(0x66, 0x66, 0x66)
COLOR_ROW_ROLE = RGBColor(0x44, 0x44, 0x44)
COLOR_ROW_DATE = RGBColor(0x55, 0x55, 0x55)
COLOR_SECTION = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_NAME_BORDER = "2C2C2C"
SHADE_SECTION = "F0F0F0"
SHADE_INTRO_BORDER = "D8D8D8"

DEFAULT_PRESET = "fill"
FIXED_LINE_HEIGHT = 1.3
# Word renders 微软雅黑 slightly taller than Chromium; scale spacing to match PDF one-page fit.
DOCX_SPACING_SCALE = 0.88
DOCX_LINE_SCALE = 0.96


def _pt(value: str) -> float:
    return float(value.rstrip("pt"))


def _cm(value: str) -> float:
    return float(value.rstrip("cm"))


def _mm(value: str) -> float:
    return float(value.rstrip("mm"))


class LayoutConfig:
    """Runtime layout values mirrored from markdown_resume_to_pdf PRESETS."""

    def __init__(self, preset_name: str = DEFAULT_PRESET) -> None:
        preset = PRESETS[preset_name]
        self.preset_name = preset_name
        self.body_size = _pt(preset["body_font_size"])
        self.name_size = _pt(preset["name_font_size"])
        self.contact_size = _pt(preset["contact_font_size"])
        self.h2_size = _pt(preset["h2_font_size"])
        self.row_heading_size = _pt(preset["row_heading_font_size"])
        self.project_intro_size = _pt(preset["project_intro_font_size"])
        self.line_spacing = float(preset["line_height"])
        self.photo_width_pt = _pt(preset["photo_width"])
        self.photo_height_pt = _pt(preset["photo_height"])
        self.margin_top_cm = _cm(preset["margin_top"])
        self.margin_bottom_cm = _cm(preset["margin_bottom"])
        self.margin_x_cm = _cm(preset["margin_x"])
        self.body_padding_top_mm = _mm(preset["body_padding_top"])
        self.h2_margin_top = _pt(preset["h2_margin_top"])
        self.section_margin_after = _pt(preset["section_margin_after"])
        self.paragraph_margin_after = _pt(preset["paragraph_margin_after"])
        self.task_margin_after = _pt(preset["task_margin_after"])
        self.project_group_margin_top = _pt(preset["project_group_margin_top"])
        self.project_group_margin_after = _pt(preset["project_group_margin_after"])
        self.project_intro_margin_after = _pt(preset["project_intro_margin_after"])
        self.company_gap_after_task = _pt(preset["company_gap_after_task"])
        self.skill_item_margin_after = _pt(preset["skill_item_margin_after"])
        self.docx_spacing_scale = DOCX_SPACING_SCALE
        self.docx_line_scale = DOCX_LINE_SCALE

    def scale_pt(self, value: float) -> float:
        return value * self.docx_spacing_scale

    def exact_line_pt(self, font_size: float | None = None, line_height: float | None = None) -> float:
        size = self.body_size if font_size is None else font_size
        height = self.line_spacing if line_height is None else line_height
        return size * height * self.docx_line_scale

    def content_width_cm(self) -> float:
        return 21.0 - 2 * self.margin_x_cm


_LAYOUT = LayoutConfig(DEFAULT_PRESET)


def hex_rgb(color: RGBColor) -> str:
    return f"{color.rgb:06X}"


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    paragraph._element.get_or_add_pPr().append(shading)


def set_paragraph_border(
    paragraph,
    side: str,
    color_hex: str,
    size: int,
    space: str = "4",
) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    element = OxmlElement(f"w:{side}")
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), str(size))
    element.set(qn("w:space"), space)
    element.set(qn("w:color"), color_hex)
    existing = p_bdr.find(qn(f"w:{side}"))
    if existing is not None:
        p_bdr.remove(existing)
    p_bdr.append(element)


def configure_run_font(
    run,
    size_pt: float | None = None,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = COLOR_BODY,
) -> None:
    if size_pt is None:
        size_pt = _LAYOUT.body_size
    run.font.name = FONT_NAME
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(
    paragraph,
    before_pt: float = 0,
    after_pt: float | None = None,
    line_spacing: float | None = None,
    font_size: float | None = None,
    fixed_line_height: float | None = None,
) -> None:
    if after_pt is None:
        after_pt = _LAYOUT.paragraph_margin_after
    before_pt = _LAYOUT.scale_pt(before_pt)
    after_pt = _LAYOUT.scale_pt(after_pt)
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if line_spacing is not None:
        exact = (font_size or _LAYOUT.body_size) * line_spacing * _LAYOUT.docx_line_scale
    elif fixed_line_height is not None:
        exact = (font_size or _LAYOUT.body_size) * fixed_line_height * _LAYOUT.docx_line_scale
    else:
        exact = _LAYOUT.exact_line_pt(font_size)
    fmt.line_spacing = Pt(exact)


def add_formatted_runs(
    paragraph,
    text: str,
    size_pt: float | None = None,
    bold_default: bool = False,
    italic_default: bool = False,
    color_default: RGBColor | None = COLOR_BODY,
) -> None:
    if size_pt is None:
        size_pt = _LAYOUT.body_size
    """Render inline **bold** markers into separate runs."""
    pattern = re.compile(r"\*\*([^*]+)\*\*")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            configure_run_font(
                run,
                size_pt=size_pt,
                bold=bold_default,
                italic=italic_default,
                color=color_default,
            )
        run = paragraph.add_run(match.group(1))
        configure_run_font(
            run,
            size_pt=size_pt,
            bold=True,
            italic=italic_default,
            color=color_default if not italic_default else color_default,
        )
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        configure_run_font(
            run,
            size_pt=size_pt,
            bold=bold_default,
            italic=italic_default,
            color=color_default,
        )
    if pos == 0 and not text:
        return


def add_section_heading(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(
        paragraph,
        before_pt=_LAYOUT.h2_margin_top,
        after_pt=_LAYOUT.section_margin_after,
        line_spacing=1.2,
        font_size=_LAYOUT.h2_size,
    )
    set_paragraph_shading(paragraph, SHADE_SECTION)
    set_paragraph_border(paragraph, "left", COLOR_NAME_BORDER, size=20, space="6")
    run = paragraph.add_run(title)
    configure_run_font(run, size_pt=_LAYOUT.h2_size, bold=True, color=COLOR_SECTION)


def add_row_heading(document: Document, company: str, middle: str, date: str, before_pt: float = 0) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before_pt=before_pt, after_pt=1, fixed_line_height=FIXED_LINE_HEIGHT, font_size=_LAYOUT.row_heading_size)
    content_width = _LAYOUT.content_width_cm()
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(content_width / 2), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Cm(content_width), WD_TAB_ALIGNMENT.RIGHT)

    run0 = paragraph.add_run(company)
    configure_run_font(run0, size_pt=_LAYOUT.row_heading_size, bold=True)
    paragraph.add_run("\t")
    run1 = paragraph.add_run(middle)
    configure_run_font(run1, size_pt=_LAYOUT.row_heading_size, color=COLOR_ROW_ROLE)
    paragraph.add_run("\t")
    run2 = paragraph.add_run(date)
    configure_run_font(run2, size_pt=_LAYOUT.row_heading_size, color=COLOR_ROW_DATE)


def add_row_heading_table(document: Document, company: str, middle: str, date: str) -> None:
    add_row_heading(document, company, middle, date)


def add_body_paragraph(
    document: Document,
    text: str,
    after_pt: float | None = None,
    indent_cm: float = 0,
    italic: bool = False,
    color: RGBColor | None = COLOR_BODY,
    size_pt: float | None = None,
) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after_pt=after_pt)
    if indent_cm:
        paragraph.paragraph_format.left_indent = Cm(indent_cm)
    add_formatted_runs(
        paragraph,
        text,
        size_pt=size_pt,
        italic_default=italic,
        color_default=color,
    )


def add_project_intro(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(
        paragraph,
        after_pt=_LAYOUT.project_intro_margin_after,
        fixed_line_height=FIXED_LINE_HEIGHT,
        font_size=_LAYOUT.project_intro_size,
    )
    paragraph.paragraph_format.left_indent = Cm(0.35)
    set_paragraph_border(paragraph, "left", SHADE_INTRO_BORDER, size=10, space="8")
    add_formatted_runs(
        paragraph,
        text,
        size_pt=_LAYOUT.project_intro_size,
        italic_default=True,
        color_default=COLOR_PROJECT_INTRO,
    )


def add_project_group(document: Document, label: str, title: str, intro: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(
        paragraph,
        before_pt=_LAYOUT.project_group_margin_top,
        after_pt=_LAYOUT.project_group_margin_after,
        fixed_line_height=FIXED_LINE_HEIGHT,
    )
    run = paragraph.add_run(f"{label}：{title}")
    configure_run_font(run, size_pt=_LAYOUT.body_size, bold=True)
    if intro:
        add_project_intro(document, intro)


def add_project_task(document: Document, prefix: str, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after_pt=_LAYOUT.task_margin_after)
    paragraph.paragraph_format.left_indent = Cm(0.35)
    if prefix:
        run_prefix = paragraph.add_run(prefix)
        configure_run_font(run_prefix, size_pt=_LAYOUT.body_size)
    add_formatted_runs(paragraph, text, size_pt=_LAYOUT.body_size)


def add_skill_item(document: Document, number: int, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after_pt=_LAYOUT.skill_item_margin_after)
    paragraph.paragraph_format.left_indent = Cm(0.35)
    # Strip leading "N. " if present in source; we bake our own number.
    stripped = re.sub(r"^\d+\.\s+", "", text.strip())
    run_num = paragraph.add_run(f"{number}. ")
    configure_run_font(run_num, size_pt=_LAYOUT.body_size)
    add_formatted_runs(paragraph, stripped, size_pt=_LAYOUT.body_size)


def _pt_to_emu(value_pt: float) -> int:
    return int(round(value_pt * 12700))


def _anchor_from_inline(inline, width_emu: int, height_emu: int) -> OxmlElement:
    """Convert an inline picture to a top-right floating anchor (matches PDF layout)."""
    anchor = OxmlElement("wp:anchor")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "0")
    anchor.set("distR", "0")
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "251658240")
    anchor.set("behindDoc", "0")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1")

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "margin")
    align = OxmlElement("wp:align")
    align.text = "right"
    position_h.append(align)
    anchor.append(position_h)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "margin")
    pos_offset = OxmlElement("wp:posOffset")
    pos_offset.text = "0"
    position_v.append(pos_offset)
    anchor.append(position_v)

    extent = OxmlElement("wp:extent")
    extent.set("cx", str(width_emu))
    extent.set("cy", str(height_emu))
    anchor.append(extent)

    effect_extent = OxmlElement("wp:effectExtent")
    effect_extent.set("l", "0")
    effect_extent.set("t", "0")
    effect_extent.set("r", "0")
    effect_extent.set("b", "0")
    anchor.append(effect_extent)

    anchor.append(OxmlElement("wp:wrapNone"))

    for child in inline:
        if child.tag == qn("wp:extent"):
            continue
        anchor.append(child)
    return anchor


def add_floating_picture(
    paragraph,
    photo_path: Path,
    width_pt: float,
    height_pt: float,
) -> None:
    """Place a headshot at the top-right margin, matching markdown_resume_to_pdf.py."""
    run = paragraph.add_run()
    run.add_picture(
        str(photo_path),
        width=Pt(width_pt),
        height=Pt(height_pt),
    )
    drawing = run._r.xpath(".//w:drawing")[0]
    inline = drawing.find(qn("wp:inline"))
    if inline is None:
        return
    width_emu = _pt_to_emu(width_pt)
    height_emu = _pt_to_emu(height_pt)
    anchor = _anchor_from_inline(inline, width_emu, height_emu)
    drawing.remove(inline)
    drawing.append(anchor)


def build_header(
    document: Document,
    name: str,
    contact_lines: list[str],
    photo_path: Path | None,
) -> None:
    if photo_path and photo_path.exists():
        photo_p = document.add_paragraph()
        set_paragraph_spacing(photo_p, before_pt=0, after_pt=0)
        photo_p.paragraph_format.space_before = Pt(0)
        photo_p.paragraph_format.space_after = Pt(0)
        add_floating_picture(
            photo_p,
            photo_path,
            _LAYOUT.photo_width_pt,
            _LAYOUT.photo_height_pt,
        )

    name_p = document.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name_p, after_pt=3, fixed_line_height=FIXED_LINE_HEIGHT, font_size=_LAYOUT.name_size)
    set_paragraph_border(name_p, "bottom", COLOR_NAME_BORDER, size=12, space="4")
    name_run = name_p.add_run(name)
    configure_run_font(name_run, size_pt=_LAYOUT.name_size, bold=True)

    for line in contact_lines:
        contact_p = document.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(
            contact_p,
            before_pt=0,
            after_pt=0,
            fixed_line_height=FIXED_LINE_HEIGHT,
            font_size=_LAYOUT.contact_size,
        )
        add_formatted_runs(
            contact_p,
            line,
            size_pt=_LAYOUT.contact_size,
            color_default=COLOR_CONTACT,
        )


def build_header_table(
    document: Document,
    name: str,
    contact_lines: list[str],
    photo_path: Path | None,
) -> None:
    build_header(document, name, contact_lines, photo_path)


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    return bool(stripped) and set(stripped.replace("|", "").replace(" ", "")) <= {"-", ":"}


def is_row_heading_line(line: str, section: str | None) -> bool:
    if section not in {"教育经历", "实习经历", "项目经历"}:
        return False
    return bool(
        re.match(
            rf"^\*\*[^*]+\*\*\s+.+(?:\d{{4}}\.\d{{2}}|{DATE_END_PATTERN})",
            line,
        )
    )


def parse_resume_markdown(markdown_text: str, base_dir: Path) -> dict:
    lines = markdown_text.splitlines()
    name = ""
    contact_lines: list[str] = []
    photo_path: Path | None = None
    sections: list[tuple[str, list[str]]] = []
    current_section: str | None = None
    current_lines: list[str] = []

    for line in lines:
        heading = re.match(r"^#\s+(.+?)\s*$", line)
        if heading:
            name = heading.group(1).strip()
            continue

        image = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", line.strip())
        if image:
            rel = image.group(2).strip()
            photo_path = (base_dir / rel).resolve()
            continue

        h2 = re.match(r"^##\s+(.+?)\s*$", line.strip())
        if h2:
            if current_section:
                sections.append((current_section, current_lines))
            current_section = h2.group(1).strip()
            current_lines = []
            continue

        if current_section is None and line.strip():
            contact_lines.append(line.strip())

        if current_section:
            current_lines.append(line)

    if current_section:
        sections.append((current_section, current_lines))

    return {
        "name": name,
        "contact_lines": contact_lines,
        "photo_path": photo_path,
        "sections": sections,
    }


def render_section(document: Document, section_name: str, section_lines: list[str]) -> None:
    add_section_heading(document, section_name)
    i = 0
    last_block_was_task = False
    while i < len(section_lines):
        raw = section_lines[i]
        line = raw.strip()

        if not line or line == "---":
            i += 1
            continue

        project_group = parse_project_group(line)
        if project_group:
            label, title, intro = project_group
            add_project_group(document, label, title, intro)
            last_block_was_task = False
            i += 1
            continue

        bold_project = parse_bold_project_header(line)
        if bold_project and section_name == "实习经历":
            title_p = document.add_paragraph()
            set_paragraph_spacing(
                title_p,
                before_pt=_LAYOUT.project_group_margin_top,
                after_pt=_LAYOUT.project_group_margin_after,
                fixed_line_height=FIXED_LINE_HEIGHT,
            )
            title_run = title_p.add_run(bold_project)
            configure_run_font(title_run, size_pt=_LAYOUT.body_size, bold=True)
            last_block_was_task = False
            i += 1
            if i < len(section_lines):
                next_line = section_lines[i].strip()
                if (
                    next_line
                    and not match_sub_item(next_line)
                    and not parse_bold_project_header(next_line)
                    and not parse_project_group(next_line)
                    and not re.match(r"^(\d+)\.\s+", next_line)
                    and not match_project_task(next_line)
                ):
                    add_project_intro(document, next_line)
                    i += 1
            continue

        sub_item = match_sub_item(line)
        if sub_item:
            add_project_task(
                document,
                f"{sub_item.group(1)}）",
                sub_item.group(2),
            )
            last_block_was_task = True
            i += 1
            continue

        bullet_task = BULLET_TASK_RE.match(line)
        if bullet_task:
            add_project_task(document, "", bullet_task.group(1))
            last_block_was_task = True
            i += 1
            continue

        project_task = match_project_task(line)
        if project_task and section_name == "实习经历":
            add_project_task(document, "", line)
            last_block_was_task = True
            i += 1
            continue

        ordered = re.match(r"^(\d+)\.\s+(.+)$", line)
        if ordered and section_name == "相关技能":
            add_skill_item(document, int(ordered.group(1)), line)
            last_block_was_task = False
            i += 1
            continue

        if is_row_heading_line(line, section_name):
            middle_trim = "，," if section_name == "教育经历" else ""
            parsed = parse_row_heading(line.replace("\u3000", " "), middle_trim)
            if parsed:
                add_row_heading(
                    document,
                    *parsed,
                    before_pt=_LAYOUT.company_gap_after_task if last_block_was_task else 0,
                )
                last_block_was_task = False
                i += 1
                continue

        add_body_paragraph(document, line)
        last_block_was_task = False
        i += 1


def build_document(parsed: dict) -> Document:
    document = Document()

    # Page setup: A4, margins aligned with PDF preset
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(_LAYOUT.margin_top_cm)
    section.bottom_margin = Cm(_LAYOUT.margin_bottom_cm)
    section.left_margin = Cm(_LAYOUT.margin_x_cm)
    section.right_margin = Cm(_LAYOUT.margin_x_cm)

    # Default document font
    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(_LAYOUT.body_size)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(_LAYOUT.exact_line_pt())

    build_header(
        document,
        parsed["name"],
        parsed["contact_lines"],
        parsed["photo_path"],
    )

    for section_name, section_lines in parsed["sections"]:
        render_section(document, section_name, section_lines)

    return document


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert resume Markdown to a styled Word document."
    )
    parser.add_argument("input", type=Path, help="Input Markdown resume path.")
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        help="Output .docx path (default: same stem as input).",
    )
    parser.add_argument(
        "--preset",
        choices=PRESETS.keys(),
        default=DEFAULT_PRESET,
        help="Layout preset mirrored from markdown_resume_to_pdf.py (default: fill).",
    )
    return parser.parse_args()


def count_docx_pages_word_com(docx_path: Path) -> int | None:
    """Return page count via Word COM on Windows, or None if unavailable."""
    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(str(docx_path.resolve()))
            pages = int(doc.ComputeStatistics(2))  # wdStatisticPages = 2
            doc.Close(False)
            return pages
        finally:
            word.Quit()
    except Exception:
        return None


def count_docx_pages_via_pdf(docx_path: Path) -> int | None:
    """Convert DOCX to PDF with Word COM and count pages."""
    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        temp_pdf = docx_path.with_suffix(".pagecheck.pdf")
        try:
            doc = word.Documents.Open(str(docx_path.resolve()))
            doc.SaveAs2(str(temp_pdf.resolve()), FileFormat=17)  # wdFormatPDF = 17
            doc.Close(False)
        finally:
            word.Quit()
        if not temp_pdf.exists():
            return None
        from markdown_resume_to_pdf import count_pdf_pages

        pages = count_pdf_pages(temp_pdf)
        temp_pdf.unlink(missing_ok=True)
        return pages
    except Exception:
        return None


def main() -> int:
    global _LAYOUT
    args = parse_args()
    _LAYOUT = LayoutConfig(args.preset)
    input_path = args.input.resolve()
    if not input_path.exists():
        print(f"Input Markdown not found: {input_path}", file=sys.stderr)
        return 1

    output_path = args.output or input_path.with_suffix(".docx")
    if not output_path.is_absolute():
        output_path = Path.cwd() / output_path
    output_path = output_path.resolve()

    markdown_text = input_path.read_text(encoding="utf-8")
    parsed = parse_resume_markdown(markdown_text, input_path.parent)
    document = build_document(parsed)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    print(f"DOCX written: {output_path}")
    print(f"Layout preset: {args.preset}")

    page_count = count_docx_pages_word_com(output_path)
    if page_count is None:
        page_count = count_docx_pages_via_pdf(output_path)
    if page_count is not None:
        print(f"Page check: {page_count} page(s).")
        if page_count != 1:
            print(
                f"Warning: exported DOCX has {page_count} pages; target is exactly 1 page.",
                file=sys.stderr,
            )
            return 2
    else:
        print("Page check: skipped (Word COM unavailable).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

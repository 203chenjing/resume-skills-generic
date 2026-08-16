#!/usr/bin/env python3
"""
Convert a Chinese resume Markdown file to a styled Word (.docx) document.

Visual style mirrors markdown_resume_to_pdf.py:
gray bar section headings with left accent, name underline,
italic project intros, photo top-right, 微软雅黑 throughout.
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Reuse parsing helpers from the PDF exporter in the same directory.
_SCRIPT_DIR = Path(__file__).resolve().parent
if str(_SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPT_DIR))

from markdown_resume_to_pdf import (  # noqa: E402
    BULLET_TASK_RE,
    DATE_END_PATTERN,
    match_project_task,
    match_sub_item,
    parse_bold_project_header,
    parse_project_group,
    parse_row_heading,
)

FONT_NAME = "微软雅黑"
COLOR_BODY = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_CONTACT = RGBColor(0x55, 0x55, 0x55)
COLOR_PROJECT_INTRO = RGBColor(0x66, 0x66, 0x66)
COLOR_ROW_ROLE = RGBColor(0x44, 0x44, 0x44)
COLOR_ROW_DATE = RGBColor(0x55, 0x55, 0x55)
COLOR_SECTION = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_NAME_BORDER = "2C2C2C"
SHADE_SECTION = "F0F0F0"
SHADE_INTRO_BORDER = "D8D8D8"

# standard preset sizes (pt)
BODY_SIZE = 9.1
NAME_SIZE = 17.0
CONTACT_SIZE = 10.0
H2_SIZE = 11.0
ROW_HEADING_SIZE = 10.8
PROJECT_INTRO_SIZE = 8.6
LINE_SPACING = 1.26
PHOTO_WIDTH_PT = 52
PHOTO_HEIGHT_PT = 76


def hex_rgb(color: RGBColor) -> str:
    return f"{color.rgb:06X}"


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    paragraph._element.get_or_add_pPr().append(shading)


def set_paragraph_border(
    paragraph,
    side: str,
    color_hex: str,
    size: int,
    space: str = "4",
) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    element = OxmlElement(f"w:{side}")
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), str(size))
    element.set(qn("w:space"), space)
    element.set(qn("w:color"), color_hex)
    existing = p_bdr.find(qn(f"w:{side}"))
    if existing is not None:
        p_bdr.remove(existing)
    p_bdr.append(element)


def configure_run_font(
    run,
    size_pt: float = BODY_SIZE,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = COLOR_BODY,
) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(
    paragraph,
    before_pt: float = 0,
    after_pt: float = 2,
    line_spacing: float = LINE_SPACING,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing


def add_formatted_runs(
    paragraph,
    text: str,
    size_pt: float = BODY_SIZE,
    bold_default: bool = False,
    italic_default: bool = False,
    color_default: RGBColor | None = COLOR_BODY,
) -> None:
    """Render inline **bold** markers into separate runs."""
    pattern = re.compile(r"\*\*([^*]+)\*\*")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            configure_run_font(
                run,
                size_pt=size_pt,
                bold=bold_default,
                italic=italic_default,
                color=color_default,
            )
        run = paragraph.add_run(match.group(1))
        configure_run_font(
            run,
            size_pt=size_pt,
            bold=True,
            italic=italic_default,
            color=color_default if not italic_default else color_default,
        )
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        configure_run_font(
            run,
            size_pt=size_pt,
            bold=bold_default,
            italic=italic_default,
            color=color_default,
        )
    if pos == 0 and not text:
        return


def add_section_heading(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before_pt=5, after_pt=2)
    set_paragraph_shading(paragraph, SHADE_SECTION)
    set_paragraph_border(paragraph, "left", COLOR_NAME_BORDER, size=20, space="6")
    run = paragraph.add_run(title)
    configure_run_font(run, size_pt=H2_SIZE, bold=True, color=COLOR_SECTION)


def add_row_heading_table(document: Document, company: str, middle: str, date: str) -> None:
    table = document.add_table(rows=1, cols=3)
    table.autofit = True
    table.allow_autofit = True
    # Remove default table style borders
    for cell in table.rows[0].cells:
        cell._tc.get_or_add_tcPr().append(OxmlElement("w:tcBorders"))

    widths = [Cm(7.0), Cm(4.5), Cm(3.5)]
    cells = table.rows[0].cells
    for cell, width in zip(cells, widths):
        cell.width = width

    p0 = cells[0].paragraphs[0]
    set_paragraph_spacing(p0, after_pt=1)
    run0 = p0.add_run(company)
    configure_run_font(run0, size_pt=ROW_HEADING_SIZE, bold=True)

    p1 = cells[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p1, after_pt=1)
    run1 = p1.add_run(middle)
    configure_run_font(run1, size_pt=ROW_HEADING_SIZE, color=COLOR_ROW_ROLE)

    p2 = cells[2].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p2, after_pt=1)
    run2 = p2.add_run(date)
    configure_run_font(run2, size_pt=ROW_HEADING_SIZE, color=COLOR_ROW_DATE)


def add_body_paragraph(
    document: Document,
    text: str,
    after_pt: float = 2,
    indent_cm: float = 0,
    italic: bool = False,
    color: RGBColor | None = COLOR_BODY,
    size_pt: float = BODY_SIZE,
) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after_pt=after_pt)
    if indent_cm:
        paragraph.paragraph_format.left_indent = Cm(indent_cm)
    add_formatted_runs(
        paragraph,
        text,
        size_pt=size_pt,
        italic_default=italic,
        color_default=color,
    )


def add_project_intro(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after_pt=2)
    paragraph.paragraph_format.left_indent = Cm(0.35)
    set_paragraph_border(paragraph, "left", SHADE_INTRO_BORDER, size=10, space="8")
    add_formatted_runs(
        paragraph,
        text,
        size_pt=PROJECT_INTRO_SIZE,
        italic_default=True,
        color_default=COLOR_PROJECT_INTRO,
    )


def add_project_group(document: Document, label: str, title: str, intro: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before_pt=4, after_pt=1.5)
    run = paragraph.add_run(f"{label}：{title}")
    configure_run_font(run, size_pt=BODY_SIZE, bold=True)
    if intro:
        add_project_intro(document, intro)


def add_project_task(document: Document, prefix: str, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after_pt=3.5)
    paragraph.paragraph_format.left_indent = Cm(0.35)
    if prefix:
        run_prefix = paragraph.add_run(prefix)
        configure_run_font(run_prefix, size_pt=BODY_SIZE)
    add_formatted_runs(paragraph, text, size_pt=BODY_SIZE)


def add_skill_item(document: Document, number: int, text: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after_pt=2)
    paragraph.paragraph_format.left_indent = Cm(0.35)
    # Strip leading "N. " if present in source; we bake our own number.
    stripped = re.sub(r"^\d+\.\s+", "", text.strip())
    run_num = paragraph.add_run(f"{number}. ")
    configure_run_font(run_num, size_pt=BODY_SIZE)
    add_formatted_runs(paragraph, stripped, size_pt=BODY_SIZE)


def build_header_table(
    document: Document,
    name: str,
    contact_lines: list[str],
    photo_path: Path | None,
) -> None:
    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    # Suppress table borders
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "nil")
            borders.append(element)
        tc_pr.append(borders)

    left_cell.width = Cm(13.5)
    right_cell.width = Cm(3.5)

    # Name with underline (bottom border on paragraph)
    name_p = left_cell.paragraphs[0]
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name_p, after_pt=3)
    set_paragraph_border(name_p, "bottom", COLOR_NAME_BORDER, size=12, space="4")
    name_run = name_p.add_run(name)
    configure_run_font(name_run, size_pt=NAME_SIZE, bold=True)
    name_run.font.letter_spacing = Pt(4)  # approximate PDF letter-spacing

    for line in contact_lines:
        contact_p = left_cell.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(contact_p, before_pt=0, after_pt=0)
        add_formatted_runs(
            contact_p,
            line,
            size_pt=CONTACT_SIZE,
            color_default=COLOR_CONTACT,
        )

    if photo_path and photo_path.exists():
        photo_p = right_cell.paragraphs[0]
        photo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(photo_p, after_pt=0)
        run = photo_p.add_run()
        run.add_picture(
            str(photo_path),
            width=Pt(PHOTO_WIDTH_PT),
            height=Pt(PHOTO_HEIGHT_PT),
        )


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    return bool(stripped) and set(stripped.replace("|", "").replace(" ", "")) <= {"-", ":"}


def is_row_heading_line(line: str, section: str | None) -> bool:
    if section not in {"教育经历", "实习经历", "项目经历"}:
        return False
    return bool(
        re.match(
            rf"^\*\*[^*]+\*\*\s+.+(?:\d{{4}}\.\d{{2}}|{DATE_END_PATTERN})",
            line,
        )
    )


def parse_resume_markdown(markdown_text: str, base_dir: Path) -> dict:
    lines = markdown_text.splitlines()
    name = ""
    contact_lines: list[str] = []
    photo_path: Path | None = None
    sections: list[tuple[str, list[str]]] = []
    current_section: str | None = None
    current_lines: list[str] = []

    for line in lines:
        heading = re.match(r"^#\s+(.+?)\s*$", line)
        if heading:
            name = heading.group(1).strip()
            continue

        image = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", line.strip())
        if image:
            rel = image.group(2).strip()
            photo_path = (base_dir / rel).resolve()
            continue

        h2 = re.match(r"^##\s+(.+?)\s*$", line.strip())
        if h2:
            if current_section:
                sections.append((current_section, current_lines))
            current_section = h2.group(1).strip()
            current_lines = []
            continue

        if current_section is None and line.strip():
            contact_lines.append(line.strip())

        if current_section:
            current_lines.append(line)

    if current_section:
        sections.append((current_section, current_lines))

    return {
        "name": name,
        "contact_lines": contact_lines,
        "photo_path": photo_path,
        "sections": sections,
    }


def render_section(document: Document, section_name: str, section_lines: list[str]) -> None:
    add_section_heading(document, section_name)
    i = 0
    while i < len(section_lines):
        raw = section_lines[i]
        line = raw.strip()

        if not line or line == "---":
            i += 1
            continue

        project_group = parse_project_group(line)
        if project_group:
            label, title, intro = project_group
            add_project_group(document, label, title, intro)
            i += 1
            continue

        bold_project = parse_bold_project_header(line)
        if bold_project and section_name == "实习经历":
            title_p = document.add_paragraph()
            set_paragraph_spacing(title_p, before_pt=4, after_pt=1.5)
            title_run = title_p.add_run(bold_project)
            configure_run_font(title_run, size_pt=BODY_SIZE, bold=True)
            i += 1
            if i < len(section_lines):
                next_line = section_lines[i].strip()
                if (
                    next_line
                    and not match_sub_item(next_line)
                    and not parse_bold_project_header(next_line)
                    and not parse_project_group(next_line)
                    and not re.match(r"^(\d+)\.\s+", next_line)
                    and not match_project_task(next_line)
                ):
                    add_project_intro(document, next_line)
                    i += 1
            continue

        sub_item = match_sub_item(line)
        if sub_item:
            add_project_task(
                document,
                f"{sub_item.group(1)}）",
                sub_item.group(2),
            )
            i += 1
            continue

        bullet_task = BULLET_TASK_RE.match(line)
        if bullet_task:
            add_project_task(document, "", bullet_task.group(1))
            i += 1
            continue

        project_task = match_project_task(line)
        if project_task and section_name == "实习经历":
            add_project_task(document, "", line)
            i += 1
            continue

        ordered = re.match(r"^(\d+)\.\s+(.+)$", line)
        if ordered and section_name == "相关技能":
            add_skill_item(document, int(ordered.group(1)), line)
            i += 1
            continue

        if is_row_heading_line(line, section_name):
            middle_trim = "，," if section_name == "教育经历" else ""
            parsed = parse_row_heading(line.replace("\u3000", " "), middle_trim)
            if parsed:
                add_row_heading_table(document, *parsed)
                i += 1
                continue

        add_body_paragraph(document, line)
        i += 1


def build_document(parsed: dict) -> Document:
    document = Document()

    # Page setup: A4, margins ~2cm top/sides
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(0.55)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Default document font
    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(BODY_SIZE)

    build_header_table(
        document,
        parsed["name"],
        parsed["contact_lines"],
        parsed["photo_path"],
    )

    for section_name, section_lines in parsed["sections"]:
        render_section(document, section_name, section_lines)

    return document


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert resume Markdown to a styled Word document."
    )
    parser.add_argument("input", type=Path, help="Input Markdown resume path.")
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        help="Output .docx path (default: same stem as input).",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    input_path = args.input.resolve()
    if not input_path.exists():
        print(f"Input Markdown not found: {input_path}", file=sys.stderr)
        return 1

    output_path = args.output or input_path.with_suffix(".docx")
    if not output_path.is_absolute():
        output_path = Path.cwd() / output_path
    output_path = output_path.resolve()

    markdown_text = input_path.read_text(encoding="utf-8")
    parsed = parse_resume_markdown(markdown_text, input_path.parent)
    document = build_document(parsed)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    print(f"DOCX written: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
